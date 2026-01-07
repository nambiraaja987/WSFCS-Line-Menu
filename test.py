import streamlit as st
import requests
import csv
import io
import zipfile
from datetime import date, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

# ==============================================================================
# CONFIGURATION & BRANDING
# ==============================================================================
st.set_page_config(page_title="WSFCS Menu Generator", layout="wide")

# ==============================================================================
# GITHUB RAW FILES
# ==============================================================================
BASE_URL = "https://raw.githubusercontent.com/nambiraaja987/WSFCS-Line-Menu/main"

CSV_URL = f"{BASE_URL}/Schools.csv"
WSFCS_LOGO_URL = f"{BASE_URL}/wsfcs.png"
CHARTWELLS_LOGO_URL = f"{BASE_URL}/Chartwells.png"

# ==============================================================================
# DISCLAIMERS
# ==============================================================================
LUNCH_DISCLAIMER = (
    "A full student lunch includes a choice of one (1) entrée supplying protein and grain, "
    "two (2) vegetable side dishes, one (1) fruit side dish, and one (1) milk. "
    "Milk choices include skim white, 1% white and skim chocolate. In order to qualify as a "
    "reimbursable meal, students must choose a minimum of three (3) components and the meal "
    "must contain ½ cup of fruit or vegetable."
)

BREAKFAST_DISCLAIMER = (
    "All students must select at least 1/2 cup of fruit with their reimbursable meal. "
    "A full student breakfast includes a choice of one (1) entrée supplying protein "
    "and/or grain, up to two (2) fruit side dishes (one (1) can be a fruit juice, "
    "and one (1) milk. Milk choices include skim white, 1% white, and skim chocolate"
)

# ==============================================================================
# CONSTANTS
# ==============================================================================
EXCLUDED_ITEMS = [
    "MAYONNAISE", "KETCHUP", "MUSTARD", "RANCH DRESSING",
    "BARBECUE SAUCE", "HOT SAUCE", "PACKET", "SYRUP",
    "MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY"
]

REPRESENTATIVE_BREAKFAST = {
    "Elementary": "ashley-magnet",
    "Middle": "clemmons-middle",
    "High": "east-forsyth"
}

ELEMENTARY_LUNCH_SLUG = "ashley-magnet"
MIDDLE_LUNCH_SLUG = "hanes-magnet"

# ==============================================================================
# HELPER FUNCTIONS
# ==============================================================================
def fetch_menu_data(slug, target_date, menu_type):
    url = (
        f"https://wsfcs.api.nutrislice.com/menu/api/weeks/school/"
        f"{slug}/menu-type/{menu_type}/{target_date:%Y/%m/%d}/?format=json"
    )
    try:
        r = requests.get(url, timeout=10)
        if r.status_code == 200:
            return r.json()
    except:
        pass
    return {}

def extract_food_items(data, target_date):
    """Extracts simple list of items (for Breakfast/Elementary)."""
    items = []
    date_str = target_date.strftime("%Y-%m-%d")
    for day in data.get("days", []):
        if day.get("date") == date_str:
            for item in day.get("menu_items", []):
                food = item.get("food")
                if food and food.get("name"):
                    name = food["name"]
                    if not any(x in name.upper() for x in EXCLUDED_ITEMS):
                        items.append(name)
            break
    return items

def extract_station_data(data, target_date, is_middle_school=False):
    """Extracts station-based data (High School / Middle School)."""
    categorized = {}
    current_station = "General Menu"
    # For HS, we might default to General Menu. For MS, we usually wait for a header.
    if not is_middle_school:
        categorized[current_station] = []
        
    date_str = target_date.strftime("%Y-%m-%d")
    
    # Blocklist for Station Names (Middle School specific)
    MS_STATION_BLOCKLIST = ["MILK", "CONDIMENT", "CONDIMENTS"]

    for day in data.get("days", []):
        if day.get("date") == date_str:
            for item in day.get("menu_items", []):
                # HEADER CHECK
                is_header = False
                header_text = ""
                if item.get('is_section_title') is True:
                    is_header = True
                    header_text = item.get('text', '')
                elif item.get('food') is None and item.get('text'):
                    is_header = True
                    header_text = item.get('text', '')

                if is_header and header_text:
                    clean_header = header_text.strip()
                    if len(clean_header) > 2: 
                        current_station = clean_header
                        # Initialize if not exists
                        if current_station not in categorized:
                            categorized[current_station] = []
                    continue 

                # FOOD CHECK
                food = item.get("food")
                if food and isinstance(food, dict) and food.get("name"):
                    food_name = food["name"]
                    # Global Exclusion check
                    if any(x in food_name.upper() for x in EXCLUDED_ITEMS):
                        continue
                        
                    # Ensure list exists
                    if current_station not in categorized:
                        categorized[current_station] = []
                        
                    if food_name not in categorized[current_station]:
                        categorized[current_station].append(food_name)
            break
            
    # Remove empty stations
    final_menu = {k: v for k, v in categorized.items() if v}

    # MIDDLE SCHOOL SPECIFIC FILTERING
    if is_middle_school:
        filtered_menu = {}
        for station, items in final_menu.items():
            # Skip if station name contains "MILK" or "CONDIMENT"
            if any(bad in station.upper() for bad in MS_STATION_BLOCKLIST):
                continue
            filtered_menu[station] = items
        return filtered_menu

    return final_menu

# --- DOCUMENT CREATION FUNCTIONS ---

def create_simple_doc(content_list, disclaimer, margin_top=2.8):
    """For Elementary Lunch and Breakfast (Simple List)."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(margin_top)
    sec.bottom_margin = Inches(1.5)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    # Disclaimer Footer
    footer = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    footer.text = disclaimer
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)

    for item in content_list:
        p = doc.add_paragraph()
        r = p.add_run(item.upper())
        r.font.name = "Times New Roman"
        r.font.size = Pt(18)
        r.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def create_middle_school_doc(categorized_data, disclaimer):
    """
    Specific Middle School Formatting:
    - Top Margin: 2.5
    - Station Header: 16pt, Bold, Underlined
    - Items: 12pt, Bold
    - No Page Breaks
    """
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(2.5)  # specific to MS
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.footer_distance = Inches(0.5)

    # Footer
    footer = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    footer.text = disclaimer
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)

    stations = list(categorized_data.keys())
    
    for index, station_name in enumerate(stations):
        # 1. STATION HEADER
        head_p = doc.add_paragraph()
        if index > 0:
            head_p.paragraph_format.space_before = Pt(18) 
        
        head_run = head_p.add_run(station_name.upper())
        head_run.font.name = "Times New Roman"
        head_run.font.size = Pt(16)
        head_run.font.bold = True
        head_run.font.underline = True
        head_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head_p.paragraph_format.space_after = Pt(6)

        # 2. FOOD ITEMS
        items = categorized_data[station_name]
        for item in items:
            p = doc.add_paragraph()
            r = p.add_run(item.upper())
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def create_high_school_doc(categorized_data, disclaimer):
    """
    Specific High School Formatting:
    - Top Margin: 2.8
    - Station Header: 24pt, Bold
    - Items: 18pt, Bold
    - Page Breaks between stations
    """
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(2.8)
    sec.bottom_margin = Inches(1.5)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.footer_distance = Inches(0.8)

    # Footer
    footer = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    footer.text = disclaimer
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)

    stations = list(categorized_data.keys())

    for index, station_name in enumerate(stations):
        # Station Header
        head_p = doc.add_paragraph()
        head_run = head_p.add_run(station_name.upper())
        head_run.font.name = "Times New Roman"
        head_run.font.size = Pt(24) 
        head_run.font.bold = True
        head_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # Items
        items = categorized_data[station_name]
        for item in items:
            p = doc.add_paragraph()
            r = p.add_run(item.upper())
            r.font.name = "Times New Roman"
            r.font.size = Pt(18)
            r.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if index < len(stations) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ==============================================================================
# WEBSITE INTERFACE
# ==============================================================================
col1, col2, col3 = st.columns([1, 2, 1])
with col1:
    st.image(WSFCS_LOGO_URL, width=150)
with col2:
    st.markdown("<h1 style='text-align: center;'>Line Menu Generator</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Select any date range.</p>", unsafe_allow_html=True)
with col3:
    st.image(CHARTWELLS_LOGO_URL, width=200)

st.markdown("---")

# ==============================================================================
# SIDEBAR
# ==============================================================================
with st.sidebar:
    st.header("⚙️ Settings")

    st.subheader("1. Select Date Range")
    start_d = st.date_input("Start Date", date.today())
    end_d = st.date_input("End Date", date.today())

    st.subheader("2. Menu Categories")
    run_breakfast = st.checkbox("All Schools - Breakfast", True)
    run_ele_lunch = st.checkbox("Elementary Lunch", True)
    run_mid_lunch = st.checkbox("Middle School Lunch", True)
    run_high_lunch = st.checkbox("High School Lunch", True)

# ==============================================================================
# MAIN LOGIC
# ==============================================================================
if st.button("🚀 Generate Menus", type="primary"):

    if start_d > end_d:
        st.error("Start Date must be before End Date.")
        st.stop()

    try:
        r = requests.get(CSV_URL, timeout=10)
        r.raise_for_status()
        schools = list(csv.DictReader(io.StringIO(r.text)))
    except Exception as e:
        st.error(f"Error loading Schools.csv: {e}")
        st.stop()

    zip_buffer = io.BytesIO()
    progress = st.progress(0)
    status = st.empty()

    dates = []
    d = start_d
    while d <= end_d:
        dates.append(d)
        d += timedelta(days=1)

    total = len(dates) * 4
    done = 0

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for d in dates:
            d_str = d.strftime("%Y-%m-%d")
            parent = f"Line Menu_{d_str}/"

            # 1. BREAKFAST
            if run_breakfast:
                status.text(f"Processing Breakfast: {d_str}")
                subfolder = f"{parent}All_School_Breakfast_Menus_{d_str}/"
                
                for level, slug in REPRESENTATIVE_BREAKFAST.items():
                    data = fetch_menu_data(slug, d, "breakfast")
                    items = extract_food_items(data, d)
                    if items:
                        # Use simple doc creator
                        doc = create_simple_doc(items, BREAKFAST_DISCLAIMER)
                        zipf.writestr(f"{subfolder}{level}_Breakfast_{d_str}.docx", doc.read())
                done += 1
                progress.progress(done / total)

            # 2. ELEMENTARY LUNCH
            if run_ele_lunch:
                status.text(f"Processing Elementary Lunch: {d_str}")
                subfolder = f"{parent}Elementary_Lunch_{d_str}/"
                
                data = fetch_menu_data(ELEMENTARY_LUNCH_SLUG, d, "lunch")
                items = extract_food_items(data, d)
                if items:
                    # Use simple doc creator
                    doc = create_simple_doc(items, LUNCH_DISCLAIMER)
                    zipf.writestr(f"{subfolder}Elementary_Lunch_{d_str}.docx", doc.read())
                done += 1
                progress.progress(done / total)

            # 3. MIDDLE LUNCH (Updated Formatting)
            if run_mid_lunch:
                status.text(f"Processing Middle Lunch: {d_str}")
                subfolder = f"{parent}Middle_School_Lunch_{d_str}/" # Note: Changed folder name to match file
                
                data = fetch_menu_data(MIDDLE_LUNCH_SLUG, d, "lunch")
                # Pass is_middle_school=True for specific filtering
                stations = extract_station_data(data, d, is_middle_school=True)
                
                if stations:
                    # Use specific Middle School doc creator
                    doc = create_middle_school_doc(stations, LUNCH_DISCLAIMER)
                    zipf.writestr(f"{subfolder}Middle_Lunch_{d_str}.docx", doc.read())
                done += 1
                progress.progress(done / total)

            # 4. HIGH SCHOOL LUNCH (Updated Filenames)
            if run_high_lunch:
                status.text(f"Processing High School Lunch: {d_str}")
                subfolder = f"{parent}High_Lunch_{d_str}/"
                
                for s in schools:
                    # Logic updated to match high_lunch.py
                    school_type = s.get("Type") or s.get("type")
                    if school_type == "HS":
                        slug = s.get("Url Name") or s.get("URL") or s.get("Url")
                        name = s.get("School Name") or s.get("SchoolName")
                        
                        if slug:
                            data = fetch_menu_data(slug, d, "lunch")
                            stations = extract_station_data(data, d, is_middle_school=False)
                            if stations:
                                # Use specific High School doc creator
                                doc = create_high_school_doc(stations, LUNCH_DISCLAIMER)
                                
                                # Updated Filename logic: replace spaces, slashes, remove dots
                                safe_name = str(name).replace(" ", "_").replace("/", "-").replace(".", "")
                                zipf.writestr(f"{subfolder}{safe_name}_Lunch.docx", doc.read())
                done += 1
                progress.progress(done / total)

    st.success("Menus Generated Successfully!")
    st.download_button(
        "📥 Download ZIP",
        zip_buffer.getvalue(),
        f"Line_Menus_{start_d}_to_{end_d}.zip",
        "application/zip"
    )
